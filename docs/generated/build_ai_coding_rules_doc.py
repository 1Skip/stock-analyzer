from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_PATH = Path(__file__).with_name("AI编码协作12条执行规则.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 32, 32)
MUTED = RGBColor(92, 92, 92)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2F3"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_in):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(w * 1440) for w in widths_in)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_in):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=None, bold=None, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", style=None, size=11, bold=False, color=INK, before=0, after=6, line_spacing=1.25):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13 if level == 2 else 12, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, size=11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    r2 = p2.add_run(body)
    set_font(r2, size=10.5, color=INK)
    add_para(doc, "", after=4)


def add_rule_table(doc, rules):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [0.45, 1.55, 2.65, 1.85]
    set_table_widths(table, widths)
    headers = ["序号", "规则", "行动口径", "避免事项"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_font(r, size=10, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])

    for idx, title, action, avoid in rules:
        cells = table.add_row().cells
        values = [str(idx), title, action, avoid]
        for i, (cell, value) in enumerate(zip(cells, values)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.18
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_font(r, size=9.2 if i in (2, 3) else 9.5, bold=(i == 1), color=INK)
    add_para(doc, "", after=6)


def add_checklist(doc, items):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [0.55, 2.0, 3.95]
    set_table_widths(table, widths)
    for cell, header in zip(table.rows[0].cells, ["检查", "检查点", "通过标准"]):
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_font(r, size=10, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for label, standard in items:
        cells = table.add_row().cells
        values = ["□", label, standard]
        for i, (cell, value) in enumerate(zip(cells, values)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.2
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_font(r, size=9.8, bold=(i == 1), color=INK)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "AI 编码协作规范"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("第 "), size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    footer.runs[-1]._r.append(fld_begin)
    footer.runs[-1]._r.append(instr)
    footer.runs[-1]._r.append(fld_end)
    set_font(footer.add_run(" 页"), size=9, color=MUTED)


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("AI 编码协作 12 条执行规则")
    set_font(r, size=24, bold=True, color=INK)

    subtitle = add_para(
        doc,
        "根据用户提供截图整理，用于约束 AI 编码助手在需求澄清、代码修改、验证交付中的执行方式。",
        size=11,
        color=MUTED,
        after=14,
    )

    add_callout(
        doc,
        "核心目标",
        "减少“猜需求、过度设计、改 A 坏 B、验证不足”的错误。每次编码任务都要先明确假设，再做最小必要修改，并在每个阶段留下可验证的检查点。",
    )

    add_heading(doc, "一、适用场景", 1)
    for item in [
        "使用 Claude Code、Codex、Cursor、Copilot Chat 等 AI 编码助手进行需求实现、缺陷修复、重构或测试补齐。",
        "多人协作项目中，需要避免 AI 越界修改、重构他人代码、或在不确定条件下强行推进。",
        "长会话、大任务、多模块改动场景中，需要控制上下文、Token、验证路径和失败记录。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        set_font(p.add_run(item), size=10.8, color=INK)

    add_heading(doc, "二、12 条规则总览", 1)
    rules = [
        (1, "先想再写", "动手前先列出目标、输入、约束、风险和关键假设；不确定就问清楚。", "不要把猜测当事实，不要在需求含糊时直接改代码。"),
        (2, "简单优先", "优先用最少代码解决当前问题，复用项目已有模式、工具和接口。", "不要为了显得完整而做超前抽象、平台化或大而全设计。"),
        (3, "外科手术式修改", "只改与目标直接相关的位置，保留他人代码和无关行为。", "不要顺手重构、格式化大片文件，或把局部修复扩成全局改造。"),
        (4, "目标驱动执行", "先定义成功标准，再执行修改；验证通过后才算完成。", "不要只因为代码写完就宣布完成，也不要跳过验证。"),
        (5, "模型只做判断", "确定性逻辑交给代码、规则、路由、测试和重试机制；LLM 负责语义判断和不确定推理。", "不要让 LLM 做路径选择、重试循环、状态机、校验规则等可编程逻辑。"),
        (6, "硬 Token 预算", "单任务控制在约 4K token，会话控制在约 30K token；大任务拆成小任务。", "不要把调试日志、全量文件和无关上下文无限塞进会话。"),
        (7, "冲突选一边", "同一处问题只采用一种方案：要么修补当前设计，要么切换新设计。", "不要混合两种架构、两套状态、两种数据来源或两种 UI 模式。"),
        (8, "先读再写", "先理解导出方、调用方、共享工具和测试，再动手修改。", "不要只看报错行就改；不要没读上下游就改公共接口。"),
        (9, "测试验证意图", "测试要说明为什么这样做，覆盖行为、边界和回归风险。", "不要只测“做了什么”，也不要写与需求无关的脆弱快照测试。"),
        (10, "每步检查点", "每个阶段总结：做了什么、验证了什么、还剩什么。", "不要长时间沉默推进，也不要把所有风险堆到最后才暴露。"),
        (11, "遵循代码库惯例", "项目用什么框架、命名、错误处理、数据流，就沿用什么。", "不要把蛇形项目改成驼峰风格，或引入项目没有的范式。"),
        (12, "失败要大声", "失败、超时、跳过、无法验证都要明确记录，并说明影响范围。", "不要在不确定时写“完成”；不要掩盖失败或吞掉异常。"),
    ]
    add_rule_table(doc, rules)

    add_heading(doc, "三、执行流程模板", 1)
    steps = [
        ("1. 任务理解", "复述目标、边界、成功标准；列出不确定点。"),
        ("2. 代码阅读", "定位入口、调用链、共享工具、现有测试和相关约定。"),
        ("3. 修改计划", "说明将改哪些文件、为什么只改这些位置、风险在哪里。"),
        ("4. 小步实现", "按最小可验证单元修改，避免把无关重构混进同一批。"),
        ("5. 验证反馈", "运行最相关测试；必要时补测试；记录失败和未覆盖风险。"),
        ("6. 交付总结", "说清楚改了什么、验证了什么、还可以继续优化什么。"),
    ]
    for label, text in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r1 = p.add_run(f"{label}：")
        set_font(r1, size=10.8, bold=True, color=DARK_BLUE)
        r2 = p.add_run(text)
        set_font(r2, size=10.8, color=INK)

    add_heading(doc, "四、任务前检查清单", 1)
    checklist = [
        ("假设已写清", "任务目标、输入条件、限制范围和不确定点已经明示。"),
        ("范围已收窄", "只改与目标相关的文件，不混入顺手重构。"),
        ("方案只选一种", "没有同时保留两套互相竞争的实现路径。"),
        ("上下游已阅读", "调用方、导出方、共享工具和现有测试已经看过。"),
        ("验证已定义", "知道要跑哪些测试或手动检查，无法验证时会说明原因。"),
        ("失败会记录", "超时、跳过、异常、缓存命中/未命中都会在结果中交代。"),
    ]
    add_checklist(doc, checklist)

    add_heading(doc, "五、可直接放进提示词的版本", 1)
    prompt = (
        "在开始编码前，请先说明你的假设、成功标准和将阅读的文件；不确定就先问。"
        "实现时优先选择最简单方案，只做与目标相关的外科手术式修改，遵循当前代码库惯例。"
        "确定性逻辑用代码、测试、路由和重试机制完成，不要交给 LLM 猜。"
        "每完成一个阶段都说明做了什么、验证了什么、还剩什么。"
        "如果失败、超时或无法验证，请明确说出来，不要把不确定结果写成完成。"
    )
    add_callout(doc, "提示词片段", prompt)

    add_heading(doc, "六、使用建议", 1)
    add_para(
        doc,
        "这 12 条规则最适合放在项目级 AI 协作规范、仓库的开发约定、或单次任务的系统提示中。若任务很小，可以只使用“先想再写、简单优先、外科手术式修改、目标驱动执行”四条；若任务跨模块或长会话，则必须启用完整 12 条。",
        size=10.8,
        after=0,
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build())
